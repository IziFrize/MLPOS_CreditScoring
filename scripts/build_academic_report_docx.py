from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RAW = ROOT / "data" / "raw"
PROCESSED = ROOT / "data" / "processed"
FIGURES = ROOT / "reports" / "figures"
REPORTS = ROOT / "reports"
OUTPUT = REPORTS / "rapport_academique_MLPOS.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
GOLD_FILL = "FFF7E0"
TOTAL_DXA = 9360


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TOTAL_DXA:
        raise ValueError(f"Table widths must total {TOTAL_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TOTAL_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text: str, bold=False, color=INK, size=9.2, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], HEADER_FILL)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=DARK_BLUE, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)
    return table


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        lead.bold = True
        lead.font.color.rgb = rgb(INK)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        p.add_run(text)


def add_heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(MUTED)


def add_figure(doc: Document, filename: str, caption: str, width=5.7) -> None:
    path = FIGURES / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        picture = p.add_run().add_picture(str(path), width=Inches(width))
        picture._inline.docPr.set("descr", caption)
        picture._inline.docPr.set("title", caption)
        add_caption(doc, caption)


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_FILL) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    r = p.add_run(f"{label} ")
    r.bold = True
    r.font.color.rgb = rgb(DARK_BLUE)
    p.add_run(text)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = "MLPOS Credit Scoring | Rapport académique"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = rgb(MUTED)
    add_page_number(section.footer.paragraphs[0])


def raw_stats() -> dict[str, pd.DataFrame]:
    names = [
        "trainperf.csv", "traindemographics.csv", "trainprevloans.csv",
        "testperf.csv", "testdemographics.csv", "testprevloans.csv", "SampleSubmission.csv",
    ]
    return {name: pd.read_csv(RAW / name) for name in names}


def main() -> None:
    REPORTS.mkdir(exist_ok=True)
    data = raw_stats()
    comparison = pd.read_csv(PROCESSED / "model_comparison.csv")
    features = json.loads((PROCESSED / "feature_columns.json").read_text(encoding="utf-8"))
    leakage = json.loads((PROCESSED / "leakage_check.json").read_text(encoding="utf-8"))
    trainperf = data["trainperf.csv"]
    target_pct = trainperf["good_bad_flag"].value_counts(normalize=True).mul(100)
    doc = Document()
    configure_document(doc)

    # Cover page: editorial-cover pattern with restrained academic styling.
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MLPOS CREDIT SCORING")
    r.bold = True
    r.font.size = Pt(27)
    r.font.color.rgb = rgb(INK)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Conception, évaluation et industrialisation d'un modèle de prédiction du défaut de paiement")
    r.italic = True
    r.font.size = Pt(15)
    r.font.color.rgb = rgb(DARK_BLUE)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Rapport académique - Data Science et MLOps").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Binôme : à compléter\nModules MAL506 / BIN501\nAudit factuel du dépôt : 31 mai 2026")
    p.paragraph_format.space_after = Pt(24)
    add_callout(doc, "Périmètre.", "Ce rapport s'appuie uniquement sur le sujet officiel, le code, les données et les artefacts effectivement présents dans le dépôt. Les limites non traitées sont indiquées explicitement.")
    doc.add_page_break()

    add_heading(doc, "Résumé exécutif", 1)
    add_body(doc, "Le projet répond au besoin de credit scoring formulé dans le sujet : estimer le risque qu'un prêt courant soit classé « Bad » à partir d'informations relatives au prêt, au client et à ses remboursements antérieurs. Les données sont hétérogènes, incomplètes et déséquilibrées. La démarche retenue conserve la table de performance comme table maître, agrège l'historique par client et construit un pipeline reproductible de préparation et de modélisation.")
    add_body(doc, f"Le pipeline produit {len(features)} variables explicatives. Cinq approches sont comparées : une baseline prédisant toujours « Good », deux régressions logistiques et deux forêts aléatoires. Sur le jeu de validation stratifié, le meilleur compromis selon le F1 de la classe « Bad » est obtenu par la régression logistique pondérée avec un seuil de 0,6 : F1 Bad = 0,4789, ROC AUC = 0,7279 et PR AUC = 0,4575.")
    add_body(doc, "L'industrialisation est opérationnelle : MLflow trace les paramètres, métriques, figures et modèles ; Docker Compose isole la base SQLite dans un volume dédié au serveur MLflow ; GitLab CI compile le code, lance les tests et exécute l'entraînement. Les limites principales sont la faible couverture démographique sur le test, l'absence de validation temporelle et la profondeur limitée des notebooks.")

    add_heading(doc, "Table des matières", 1)
    add_bullets(doc, [
        "1. Introduction",
        "2. Contexte métier",
        "3. Présentation des données",
        "4. Exploration et qualité des données",
        "5. Préparation et nettoyage",
        "6. Feature engineering et justification métier",
        "7. Modélisation",
        "8. Résultats et sélection du modèle",
        "9. Tracking des expériences avec MLflow",
        "10. Docker et CI/CD",
        "11. Limites",
        "12. Perspectives",
        "13. Conclusion",
        "Annexe A. Audit de conformité à la grille",
    ])

    add_heading(doc, "1. Introduction", 1)
    add_body(doc, "Les institutions financières arbitrent en permanence entre deux risques : accorder un prêt à un emprunteur susceptible de faire défaut, ou refuser un prêt à un client solvable. Un système de credit scoring vise à objectiver cette décision au moyen d'un score de risque calculé à partir des informations disponibles au moment de la demande.")
    add_body(doc, "Le sujet demande de couvrir le cycle de vie complet d'un modèle de Machine Learning : compréhension métier, préparation de données multi-sources, modélisation, évaluation adaptée au déséquilibre des classes, traçabilité avec MLflow et industrialisation avec Docker et GitLab CI/CD. Le dépôt analysé suit cette logique au moyen de scripts Python reproductibles complétés par des notebooks d'illustration.")

    add_heading(doc, "2. Contexte métier", 1)
    add_body(doc, "La variable cible `good_bad_flag` distingue les prêts « Good » des prêts « Bad ». Dans le code, elle est convertie en `target_bad`, avec `Bad = 1` et `Good = 0`. Cette convention place volontairement la classe risquée au centre de l'analyse : rappel, précision et F1 sont calculés pour les défauts.")
    add_body(doc, "Dans un contexte bancaire, le coût d'un faux négatif est généralement significatif : un client réellement risqué est traité comme un bon dossier. Le coût d'un faux positif existe également : un client solvable peut être refusé ou faire l'objet de conditions excessives. La sélection finale doit donc rechercher un compromis explicite et non maximiser mécaniquement l'accuracy.")

    add_heading(doc, "3. Présentation des données", 1)
    add_body(doc, "Les données sont issues d'une compétition Zindi et sont réparties en fichiers train et test. Trois familles d'information sont nécessaires : la performance du prêt courant, les informations démographiques et l'historique des prêts antérieurs. `customerid` constitue la clé de rapprochement commune.")
    dataset_rows = [
        ["trainperf.csv", "4 368", "10", "Table maître train ; contient la cible"],
        ["traindemographics.csv", "4 346", "9", "Démographie train"],
        ["trainprevloans.csv", "18 183", "12", "Prêts antérieurs train"],
        ["testperf.csv", "1 450", "9", "Table maître test"],
        ["testdemographics.csv", "1 487", "9", "Démographie test"],
        ["testprevloans.csv", "5 907", "12", "Prêts antérieurs test"],
        ["SampleSubmission.csv", "1 450", "2", "Format attendu de prédiction"],
    ]
    add_table(doc, ["Fichier", "Lignes", "Colonnes", "Rôle"], dataset_rows, [2500, 900, 900, 5060])
    add_caption(doc, "Tableau 1 - Inventaire des fichiers bruts effectivement utilisés.")
    add_heading(doc, "3.1 Pourquoi utiliser CustomerID ?", 2)
    add_body(doc, "`customerid` représente l'entité métier stable : l'emprunteur. Une ligne de `trainperf.csv` ou `testperf.csv` correspond au prêt courant à scorer, tandis que plusieurs lignes de `prevloans` peuvent appartenir au même client. L'agrégation par `customerid` transforme l'historique variable en un vecteur de caractéristiques compatible avec un apprentissage supervisé.")
    add_body(doc, "Les jointures sont réalisées en `left join` à partir des tables de performance. Ce choix est essentiel : l'absence d'information démographique ou historique ne doit pas supprimer un dossier bancaire. Elle devient au contraire une information exploitable grâce aux indicateurs de présence.")

    add_heading(doc, "4. Exploration et qualité des données", 1)
    add_heading(doc, "4.1 Déséquilibre de classes", 2)
    add_body(doc, f"La cible est déséquilibrée : `Good` représente {target_pct['Good']:.2f} % des 4 368 dossiers d'entraînement, contre {target_pct['Bad']:.2f} % pour `Bad`. Ce déséquilibre explique pourquoi une baseline prédisant toujours la classe majoritaire atteint une accuracy élevée tout en étant inutilisable pour détecter le risque.")
    add_heading(doc, "4.2 Valeurs manquantes et doublons", 2)
    quality_rows = [
        ["traindemographics", "bank_branch_clients", "98,83 %", "Indicateur de présence puis suppression de la colonne brute"],
        ["traindemographics", "level_of_education_clients", "86,49 %", "Catégorie `Unknown` et indicateur de présence"],
        ["traindemographics", "employment_status_clients", "14,91 %", "Catégorie `Unknown` et indicateur de présence"],
        ["trainperf", "referredby", "86,56 %", "Remplacé par `has_referrer`"],
        ["trainprevloans", "referredby", "94,36 %", "Agrégé en ratio de présence"],
        ["testperf", "dates invalides", "853 / 859", "Dates exclues des features du prêt courant"],
    ]
    add_table(doc, ["Source", "Champ", "Constat", "Traitement"], quality_rows, [1700, 2100, 1200, 4360])
    add_caption(doc, "Tableau 2 - Principaux problèmes de qualité observés et traitement réellement implémenté.")
    add_body(doc, "`traindemographics.csv` contient 12 doublons exacts et 12 identifiants client dupliqués ; `testdemographics.csv` contient 3 doublons exacts et 3 identifiants dupliqués. La fonction `clean_demographics` supprime les doublons puis conserve une seule ligne par `customerid`. Les fichiers de performance ne comportent pas de doublon client.")
    add_heading(doc, "4.3 Couverture des sources", 2)
    coverage_rows = [
        ["Train", "4 368", "3 269 (74,84 %)", "4 359 (99,79 %)"],
        ["Test", "1 450", "385 (26,55 %)", "1 442 (99,45 %)"],
    ]
    add_table(doc, ["Population", "Dossiers", "Avec démographie", "Avec historique"], coverage_rows, [1700, 1600, 3000, 3060])
    add_caption(doc, "Tableau 3 - Couverture des sources après rapprochement par client.")
    add_callout(doc, "Point de vigilance.", "La couverture démographique chute fortement entre train et test : 74,84 % contre 26,55 %. Le modèle peut donc s'appuyer moins souvent sur ces variables au moment de la prédiction. Cette dérive doit être mesurée et surveillée avant une utilisation opérationnelle.", fill=GOLD_FILL)

    add_heading(doc, "5. Préparation et nettoyage", 1)
    add_body(doc, "La préparation est centralisée dans `src/preprocessing.py` et encapsulée dans un pipeline sklearn. Cette organisation réduit le risque de divergence entre entraînement et prédiction.")
    prep_rows = [
        ["Dates", "Conversion avec `pd.to_datetime(..., errors='coerce')`", "Neutraliser les formats invalides sans interrompre le pipeline"],
        ["Démographie", "Doublons supprimés ; une ligne conservée par client", "Garantir une cardinalité compatible avec la jointure"],
        ["Catégories", "Valeurs manquantes remplacées par `Unknown`", "Préserver l'information d'absence et éviter les pertes de lignes"],
        ["Numérique", "Imputation médiane", "Réduire l'influence des valeurs extrêmes"],
        ["Catégoriel", "Imputation par mode puis One-Hot Encoding", "Fournir des variables exploitables par les modèles sklearn"],
        ["Nouveaux niveaux", "`handle_unknown='ignore'`", "Accepter des catégories inconnues lors du scoring"],
        ["Échelle", "`StandardScaler` sur les numériques", "Rendre les coefficients logistiques comparables et stabiliser l'optimisation"],
    ]
    add_table(doc, ["Étape", "Implémentation", "Justification"], prep_rows, [1500, 3300, 4560])
    add_caption(doc, "Tableau 4 - Pipeline de préparation reproductible.")
    add_body(doc, "Pour le prêt courant, les dates sont volontairement retirées après conversion car les dates de `testperf.csv` sont largement invalides : 853 dates d'approbation et 859 dates de création ne sont pas interprétables. Cette décision limite la richesse temporelle du modèle mais évite une variable instable entre entraînement et scoring.")

    add_heading(doc, "6. Feature engineering et justification métier", 1)
    add_body(doc, f"Le fichier `feature_columns.json` recense {len(features)} variables finales. Elles se répartissent en quatre ensembles complémentaires : prêt courant, disponibilité des données, démographie et historique des remboursements.")
    feature_rows = [
        ["Prêt courant", "`loanamount`, `totaldue`, `termdays`, `loannumber`", "Mesurer l'exposition financière, l'échéance et la maturité de la relation"],
        ["Coût du prêt", "`loan_interest`, `loan_interest_rate`", "Représenter l'effort financier associé au financement"],
        ["Disponibilité", "`has_referrer`, `has_demographics`, `has_prevloans`", "Distinguer absence de donnée et valeur économique observée"],
        ["Démographie", "`age`, banque, type de compte, emploi, éducation, GPS", "Décrire le profil déclaré sans en faire une causalité automatique"],
        ["Historique volume", "`prev_nb_loans`, montants totaux/moyens/min/max/std", "Capturer l'expérience de crédit et l'évolution de l'exposition"],
        ["Historique durée", "durées moyenne/min/max et remboursement moyen", "Décrire le rythme de remboursement"],
        ["Historique retard", "ratios de retard/avance et jours de retard moyen/min/max", "Mesurer directement le comportement passé de remboursement"],
        ["Dernier prêt", "montant, total dû, durée et retard du dernier prêt", "Accorder une visibilité au comportement le plus récent"],
    ]
    add_table(doc, ["Groupe", "Variables", "Intérêt métier"], feature_rows, [1550, 3500, 4310])
    add_caption(doc, "Tableau 5 - Groupes de variables construites et justification métier.")
    add_heading(doc, "6.1 Justification métier des variables", 2)
    add_body(doc, "Les variables de retard sont les plus directement reliées au risque : un client ayant remboursé après l'échéance a démontré une difficulté ou une moindre discipline de paiement. Les ratios `prev_late_ratio` et `prev_early_ratio` résument la régularité du comportement sur plusieurs prêts, tandis que `prev_last_days_late` donne davantage de visibilité à la situation récente.")
    add_body(doc, "Les montants et durées décrivent l'exposition. Un montant plus élevé, un total dû important ou une durée plus longue peuvent augmenter la contrainte de remboursement. Les variables de dispersion (`prev_std_loanamount`) signalent une trajectoire moins stable. Le nombre de prêts antérieurs reflète l'historique disponible et la maturité de la relation.")
    add_body(doc, "Les variables démographiques peuvent améliorer la discrimination mais doivent être interprétées avec prudence. Elles sont incomplètes, leur couverture varie fortement entre train et test et certaines dimensions peuvent soulever des enjeux d'équité. Le projet les conserve avec des indicateurs de disponibilité ; une analyse de biais reste nécessaire avant tout usage réel.")
    add_heading(doc, "6.2 Contrôle anti-fuite", 2)
    add_body(doc, f"Le contrôle implémenté compare les prêts historiques et le prêt courant par numéro et par date d'approbation. Sur {leakage['rows_checked']:,} comparaisons, `loannumber_leaks = {leakage['loannumber_leaks']}` et `approveddate_leaks = {leakage['approveddate_leaks']}`. Aucun prêt futur n'est identifié dans l'historique utilisé.".replace(",", " "))

    add_heading(doc, "7. Modélisation", 1)
    add_body(doc, "Le dataset est séparé en apprentissage et validation selon un split stratifié `80/20` (`test_size = 0.20`, `random_state = 42`). La stratification conserve la proportion de dossiers `Bad`. Quatre seuils de décision sont testés : 0,3 ; 0,4 ; 0,5 ; 0,6. Pour chaque modèle, le seuil maximisant d'abord le F1 Bad puis le rappel Bad est retenu.")
    model_rows = [
        ["Baseline naïve", "`DummyClassifier(constant=0)`", "Toujours `Good`", "Référence minimale ; ne détecte aucun défaut"],
        ["Régression logistique", "`max_iter=1000`, `random_state=42`", "Simple, stable, interprétable", "Relations essentiellement linéaires"],
        ["Régression logistique pondérée", "`class_weight='balanced'`", "Rééquilibre l'apprentissage", "Interprétation à compléter par l'étude des coefficients"],
        ["Random Forest simple", "`n_estimators=200`, `max_depth=8`, `min_samples_leaf=2`", "Non-linéarités et interactions", "Risque d'overfitting supérieur"],
        ["Random Forest pondérée", "`n_estimators=200`, `max_depth=8`, `min_samples_leaf=5`, `class_weight='balanced'`", "Non-linéarités avec priorité accrue à la classe minoritaire", "Écart train-validation observé"],
    ]
    add_table(doc, ["Modèle", "Hyperparamètres clés", "Atout", "Limite"], model_rows, [1800, 2850, 2200, 2510], font_size=8.6)
    add_caption(doc, "Tableau 6 - Modèles réellement entraînés et hyperparamètres clés.")
    add_body(doc, "La baseline est indispensable : elle démontre qu'une accuracy correcte peut masquer une absence totale de détection des dossiers risqués. La régression logistique fournit une référence robuste et explicable. Les forêts aléatoires testent ensuite l'apport de relations non linéaires et d'interactions plus complexes.")

    add_heading(doc, "8. Résultats et sélection du modèle", 1)
    result_rows = []
    for row in comparison.itertuples():
        result_rows.append([
            row.run_name.replace("_", " "),
            f"{row.threshold:.1f}",
            f"{row.accuracy:.4f}",
            f"{row.precision_bad:.4f}",
            f"{row.recall_bad:.4f}",
            f"{row.f1_bad:.4f}",
            f"{row.roc_auc:.4f}",
            f"{row.pr_auc:.4f}",
        ])
    add_table(doc, ["Modèle", "Seuil", "Accuracy", "Préc. Bad", "Rappel Bad", "F1 Bad", "ROC AUC", "PR AUC"], result_rows, [3200, 650, 950, 950, 950, 900, 900, 860], font_size=8.4)
    add_caption(doc, "Tableau 7 - Comparaison réelle des performances sur le jeu de validation.")
    add_heading(doc, "8.1 Pourquoi l'accuracy ne suffit pas", 2)
    add_body(doc, "La baseline naïve atteint l'accuracy la plus élevée, 0,7826, parce qu'elle prédit systématiquement la classe majoritaire `Good`. Elle obtient pourtant une précision Bad, un rappel Bad et un F1 Bad nuls. Elle ne répond donc pas au besoin métier. La PR AUC de 0,2174 correspond au niveau de prévalence de la classe positive et confirme l'absence de pouvoir discriminant.")
    add_heading(doc, "8.2 Choix du modèle final", 2)
    add_body(doc, "Le modèle retenu par le pipeline est `logistic_regression_balanced` avec un seuil de 0,6. Il obtient le meilleur F1 Bad, 0,4789, à égalité entre précision et rappel. Son ROC AUC de 0,7279 et sa PR AUC de 0,4575 sont également les meilleures valeurs du tableau. Sur 190 dossiers `Bad` de validation, 91 sont correctement identifiés ; 99 sont manqués. Le modèle produit également 99 faux positifs.")
    add_callout(doc, "Interprétation métier.", "Le modèle final n'est pas un outil de décision automatique. Il constitue un premier score de risque reproductible. Le seuil devra être recalibré selon les coûts réels d'un défaut non détecté et d'un dossier solvable refusé.")
    add_figure(doc, "logistic_regression_balanced_confusion_matrix.png", "Figure 1 - Matrice de confusion du modèle final : régression logistique pondérée, seuil 0,6.")
    add_figure(doc, "logistic_regression_balanced_roc_curve.png", "Figure 2 - Courbe ROC du modèle final.")
    add_figure(doc, "logistic_regression_balanced_precision_recall_curve.png", "Figure 3 - Courbe précision-rappel du modèle final.")
    add_heading(doc, "8.3 Risque d'overfitting", 2)
    overfit_rows = []
    for row in comparison.itertuples():
        overfit_rows.append([
            row.run_name.replace("_", " "),
            f"{row.train_f1_bad:.4f}",
            f"{row.f1_bad:.4f}",
            f"{row.train_f1_bad - row.f1_bad:+.4f}",
        ])
    add_table(doc, ["Modèle", "F1 Bad train", "F1 Bad validation", "Écart"], overfit_rows, [3600, 1900, 2100, 1760])
    add_caption(doc, "Tableau 8 - Comparaison train-validation du F1 Bad.")
    add_body(doc, "Les régressions logistiques ne présentent pas de signal d'overfitting : leur F1 Bad de validation est légèrement supérieur au score d'entraînement. Les Random Forest présentent au contraire un écart d'environ 0,16 point de F1 Bad entre train et validation. Leur capacité plus élevée n'améliore pas la généralisation dans cette configuration.")
    add_figure(doc, "random_forest_balanced_feature_importance.png", "Figure 4 - Importance des variables pour la Random Forest pondérée. Cette figure sert à l'interprétation du modèle avancé, sans remplacer une analyse causale.")

    add_heading(doc, "9. Tracking des expériences avec MLflow", 1)
    add_body(doc, "MLflow est intégré dans `src/train.py`. L'expérience `credit_scoring_mlpos` regroupe les runs. Pour chaque modèle, le script ouvre un run nommé, logge les paramètres, les métriques de validation et d'entraînement, les figures, les rapports de classification, les colonnes finales et le pipeline sklearn.")
    mlflow_rows = [
        ["Paramètres", "Nom du modèle, taille du test, seed, lignes train/validation, nombre de variables, mapping cible et hyperparamètres clés"],
        ["Métriques", "Seuil, accuracy, précision Bad, rappel Bad, F1 Bad, F1 macro, ROC AUC, PR AUC et variantes train"],
        ["Artefacts", "Figures, rapports de classification, `feature_columns.json`"],
        ["Modèles", "Pipeline sklearn enregistré avec `mlflow.sklearn.log_model`"],
        ["Persistance Compose", "Base SQLite et artefacts isolés dans deux volumes Docker nommés"],
    ]
    add_table(doc, ["Élément MLflow", "Contenu réellement enregistré"], mlflow_rows, [2200, 7160])
    add_caption(doc, "Tableau 9 - Traçabilité MLflow implémentée.")
    add_body(doc, "La stack Docker Compose a été vérifiée après correction : le serveur MLflow est `healthy`, SQLite retourne `integrity ok`, l'entraînement se termine avec le code `0` et les artefacts sont servis par le tracking server. L'isolation de SQLite évite que le client d'entraînement et l'UI ne tentent de migrer simultanément la même base.")

    add_heading(doc, "10. Docker et CI/CD", 1)
    add_heading(doc, "10.1 Reproductibilité Docker", 2)
    add_body(doc, "Le `Dockerfile` part de `python:3.10-slim`, installe `requirements.txt`, copie le projet et lance `python src/train.py`. Le fichier `docker-compose.yml` déclare deux services : le serveur MLflow et l'entraînement. Le service d'entraînement reçoit `MLFLOW_TRACKING_URI=http://mlflow-ui:5000` et attend le healthcheck du serveur.")
    add_bullets(doc, [
        "Installation locale : `pip install -r requirements.txt`.",
        "Tests locaux : `python -m pytest tests -q`.",
        "Entraînement local : `python src/train.py`.",
        "Prédiction locale : `python src/predict.py`.",
        "Stack MLOps : `docker compose up --build`.",
        "Interface MLflow : `http://localhost:5000`.",
    ])
    add_heading(doc, "10.2 Pipeline GitLab CI", 2)
    ci_rows = [
        ["test", "Compilation Python puis `pytest tests/`", "Détecter erreurs syntaxiques et régressions couvertes"],
        ["train", "`python src/train.py`", "Reproduire l'entraînement"],
        ["artefacts", "`models/`, `reports/figures/`, `data/processed/`", "Conserver les sorties pendant une semaine"],
    ]
    add_table(doc, ["Étape", "Commande / contenu", "Objectif"], ci_rows, [1600, 4000, 3760])
    add_caption(doc, "Tableau 10 - Pipeline CI présent dans `.gitlab-ci.yml`.")
    add_body(doc, "La CI est cohérente avec un projet académique minimal. Elle ne contient toutefois pas d'étape `validate` autonome exécutant `src/predict.py` et contrôlant la soumission générée. L'existence du fichier CI est vérifiée dans le dépôt ; son exécution effective sur une instance GitLab ne peut pas être déduite des fichiers locaux.")

    add_heading(doc, "11. Limites", 1)
    add_bullets(doc, [
        "Couverture démographique instable : 74,84 % des dossiers train contre 26,55 % des dossiers test.",
        "Variables fortement incomplètes : agence bancaire, éducation et référent.",
        "Dates de performance test largement invalides, empêchant leur exploitation directe.",
        "Validation aléatoire stratifiée mais non temporelle ; le comportement hors période n'est pas mesuré.",
        "F1 Bad du meilleur modèle inférieur à 0,5 : le score constitue une première base, pas un système prêt pour une décision automatisée.",
        "Notebooks présents mais courts, sans sorties enregistrées ni EDA graphique approfondie.",
        "Analyse d'équité absente malgré l'utilisation de variables démographiques.",
        "SQLite adapté à cette stack mono-serveur, mais PostgreSQL préférable pour une plateforme partagée ou de production.",
    ])

    add_heading(doc, "12. Perspectives", 1)
    perspective_rows = [
        ["Court terme", "Enrichir les notebooks exécutés", "Rendre visibles distributions, valeurs manquantes, jointures et résultats"],
        ["Court terme", "Ajouter une étape CI `validate`", "Exécuter `src/predict.py` et tester la conformité de `submission.csv`"],
        ["Court terme", "Étudier une validation temporelle", "Mesurer la robustesse sur une période postérieure"],
        ["Modélisation", "Tester XGBoost ou LightGBM", "Comparer une méthode de boosting adaptée aux interactions"],
        ["Optimisation", "Recherche d'hyperparamètres et calibration", "Optimiser le compromis métier et la qualité des probabilités"],
        ["Interprétation", "SHAP et analyse de stabilité", "Expliquer les décisions et surveiller la dérive"],
        ["Équité", "Mesurer les biais par sous-groupes", "Vérifier l'absence d'effets indésirables sur les profils clients"],
        ["Déploiement", "API FastAPI et monitoring", "Exposer le scoring, tracer les prédictions et détecter la dérive"],
        ["Industrialisation", "Migrer SQLite vers PostgreSQL si nécessaire", "Supporter plusieurs utilisateurs et une charge accrue"],
    ]
    add_table(doc, ["Horizon", "Amélioration", "Bénéfice attendu"], perspective_rows, [1500, 3300, 4560])
    add_caption(doc, "Tableau 11 - Améliorations réalistes, sans exagérer la maturité actuelle.")

    add_heading(doc, "13. Conclusion", 1)
    add_body(doc, "Le dépôt MLPOS Credit Scoring fournit une chaîne cohérente et reproductible pour un projet universitaire de Data Science et MLOps. Les trois sources sont nettoyées et fusionnées sans perte de dossiers maîtres, l'historique est agrégé en variables métier pertinentes, cinq approches sont comparées avec des métriques adaptées au déséquilibre et le meilleur pipeline est sérialisé.")
    add_body(doc, "La régression logistique pondérée constitue le meilleur compromis actuel. Elle surpasse les alternatives sur le F1 Bad, la ROC AUC et la PR AUC tout en limitant l'overfitting observé sur les Random Forest. MLflow, Docker Compose et GitLab CI apportent la traçabilité et la reproductibilité attendues. Les perspectives prioritaires sont l'enrichissement de l'EDA, la validation temporelle, l'analyse de biais et le renforcement de la CI.")

    doc.add_page_break()
    add_heading(doc, "Annexe A. Audit de conformité à la grille", 1)
    audit_rows = [
        ["Compréhension métier et données", "Conforme", "Contexte, sources, cible, clé de jointure et déséquilibre explicités"],
        ["Nettoyage et qualité", "Conforme", "Doublons, manquants, types, encodage, standardisation et anomalies de dates traités"],
        ["Feature engineering", "Conforme", "43 variables finales ; agrégations historiques et contrôle anti-fuite"],
        ["Modélisation et performances", "Conforme", "Baseline, modèles simples/avancés, seuils et métriques adaptées"],
        ["MLflow", "Conforme", "Paramètres, métriques, figures et modèles enregistrés"],
        ["Docker", "Conforme et vérifié", "Stack Compose testée ; serveur healthy ; entraînement code 0"],
        ["CI/CD", "Partiel", "Tests et entraînement présents ; étape autonome de validation à ajouter"],
        ["Rapport et interprétation", "Conforme", "Analyse factuelle, limites honnêtes et perspectives réalistes"],
        ["Bonus", "Absent", "Biais, SHAP avancé, API et monitoring proposés en perspectives"],
    ]
    add_table(doc, ["Critère", "État", "Éléments vérifiés"], audit_rows, [2450, 1700, 5210])
    add_caption(doc, "Tableau 12 - Vérification finale par rapport à la grille du sujet officiel.")
    add_heading(doc, "Annexe B. Points faibles restants à corriger sans refonte", 1)
    add_bullets(doc, [
        "Exécuter les notebooks et conserver les sorties utiles pour rendre l'EDA immédiatement lisible.",
        "Ajouter à GitLab CI une étape de validation lançant la prédiction et testant la structure de la soumission.",
        "Documenter le coût métier des faux positifs et faux négatifs pour choisir le seuil de décision.",
        "Comparer les performances avec et sans variables démographiques afin de mesurer leur apport réel et le risque de dérive.",
        "Ajouter une validation temporelle et une première analyse d'équité avant toute prétention opérationnelle.",
    ])
    add_heading(doc, "Annexe C. Sources internes utilisées", 1)
    add_bullets(doc, [
        "Sujet officiel : `PROJET MLPOS Cours MAL506BIN501.docx`.",
        "Code : `src/preprocessing.py`, `src/features.py`, `src/train.py`, `src/evaluate.py`, `src/predict.py`.",
        "MLOps : `Dockerfile`, `docker-compose.yml`, `.gitlab-ci.yml`, `README.md`.",
        "Résultats : `data/processed/model_comparison.csv`, rapports de classification, figures et `leakage_check.json`.",
        "Vérifications : tests pytest, exécution Docker Compose et contrôle d'intégrité SQLite.",
    ])
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
