from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PROCESSED = ROOT / "data" / "processed"
REPORTS = ROOT / "reports"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table_from_dataframe(doc: Document, df: pd.DataFrame, title: str, max_rows: int = 10) -> None:
    add_heading(doc, title, level=2)
    table_df = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, col in enumerate(table_df.columns):
        header_cells[idx].text = str(col)
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                value = round(value, 4)
            cells[idx].text = str(value)
    doc.add_paragraph()


def main() -> None:
    REPORTS.mkdir(exist_ok=True)
    comparison_path = PROCESSED / "model_comparison.csv"
    leakage_path = PROCESSED / "leakage_check.json"

    comparison = pd.read_csv(comparison_path) if comparison_path.exists() else pd.DataFrame()
    leakage_text = leakage_path.read_text(encoding="utf-8") if leakage_path.exists() else "Controle non disponible."

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MLPOS Credit Scoring")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Projet Machine Learning / MLOps pour la prediction du risque de defaut").italic = True

    add_heading(doc, "1. Introduction", 1)
    doc.add_paragraph(
        "Ce projet construit un pipeline de credit scoring pour predire si un client sera Good ou Bad. "
        "L'objectif metier est d'identifier les profils a risque sans se limiter a l'accuracy, car la classe Bad est minoritaire."
    )

    add_heading(doc, "2. Presentation des donnees", 1)
    add_bullets(
        doc,
        [
            "trainperf.csv : table principale d'entrainement avec la target good_bad_flag.",
            "traindemographics.csv : informations demographiques, incompletes et dedupliquees par customerid.",
            "trainprevloans.csv : historique des anciens prets, agrege par customerid.",
            "testperf.csv : table principale de test sans target.",
            "SampleSubmission.csv : format final attendu avec customerid et Good_Bad_flag.",
        ],
    )

    add_heading(doc, "3. Exploration et nettoyage", 1)
    add_bullets(
        doc,
        [
            "Les jointures utilisent toujours trainperf ou testperf comme table maitre.",
            "Les jointures sont des left join pour ne pas perdre de lignes.",
            "Les demographics sont dedupliques par customerid.",
            "Les dates de testperf sont considerees non fiables et ne sont pas utilisees comme features principales.",
            "La target est encodee avec Bad = 1 et Good = 0.",
        ],
    )

    add_heading(doc, "4. Feature engineering", 1)
    doc.add_paragraph(
        "Les anciens prets sont agreges par client. Les features couvrent les montants, durees, interets, retards, "
        "remboursements anticipes et informations du dernier pret precedent."
    )
    add_bullets(
        doc,
        [
            "Features du pret courant : loannumber, loanamount, totaldue, termdays, has_referrer, loan_interest, loan_interest_rate.",
            "Features demographics : age, banque, type de compte, statut professionnel, niveau d'etudes, GPS et flags de presence.",
            "Features prevloans : nombre de prets, montants moyens et totaux, retards, durees, taux d'interet et dernier pret.",
        ],
    )

    add_heading(doc, "5. Controle anti-fuite de donnees", 1)
    doc.add_paragraph(
        "Le projet verifie que les prets precedents ne contredisent pas le pret courant via loannumber et les dates disponibles."
    )
    doc.add_paragraph(leakage_text)

    add_heading(doc, "6. Modelisation", 1)
    add_bullets(
        doc,
        [
            "Baseline naive : predit toujours Good.",
            "LogisticRegression : modele simple et interpretable.",
            "LogisticRegression balanced : prise en compte du desequilibre de classes.",
            "RandomForestClassifier : modele non lineaire, teste en version simple et balanced.",
            "Split train/validation stratifie avec test_size = 0.2 et random_state = 42.",
        ],
    )

    if not comparison.empty:
        add_table_from_dataframe(doc, comparison, "Comparaison des modeles", max_rows=8)

    add_heading(doc, "7. MLOps et reproductibilite", 1)
    add_bullets(
        doc,
        [
            "MLflow logge les parametres, metriques, figures, rapports et modeles.",
            "Dockerfile et docker-compose.yml permettent d'executer l'entrainement et l'UI MLflow.",
            "GitLab CI compile le code, lance les tests et execute l'entrainement.",
            "Les tests pytest verifient preprocessing, features, schema et format de soumission.",
        ],
    )

    add_heading(doc, "8. Limites", 1)
    add_bullets(
        doc,
        [
            "Les donnees demographiques couvrent peu le test.",
            "Certaines colonnes sont tres incompletes.",
            "Les dates de testperf sont mal formatees.",
            "La validation est aleatoire et non temporelle.",
            "Le seuil final doit etre justifie selon le compromis metier precision/recall sur la classe Bad.",
        ],
    )

    add_heading(doc, "9. Conclusion", 1)
    if not comparison.empty:
        best = comparison.iloc[0]
        doc.add_paragraph(
            f"Le meilleur run actuel est {best['run_name']} avec un seuil {best['threshold']} "
            f"et un F1 Bad de {best['f1_bad']:.4f}."
        )
    doc.add_paragraph(
        "Le projet est volontairement simple, reproductible et adapte a un rendu universitaire Master."
    )

    output = REPORTS / "rapport_MLPOS.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
